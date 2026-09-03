"""
Getting words out of a file.

Three routes, and which one a file takes decides whether it can ever be found by
its contents:

    text      PDF, DOCX, and anything already text — read directly
    ocr       images, and PDFs whose pages hold no text layer
    skipped   video, archives, executables, and anything too large

The OCR route is the one that matters for the failure this was built for. An
Aadhaar card saved as `scan001.jpg` has nothing useful in its name and nothing
in its metadata; the only thing that says what it is, is the word AADHAAR
printed on the image. It is read **at index time and never at query time** — a
search that OCR'd on demand would take minutes and would do it again for the
next search.

**Windows' own OCR, not Tesseract.** `Windows.Media.Ocr` has shipped since
Windows 10 and is already on this machine — checked before choosing it, along
with pytesseract, which is installed but whose `tesseract.exe` is not, and would
therefore have meant asking the owner to install a second thing before search
worked.
"""
from __future__ import annotations

import logging
import os
import subprocess
from pathlib import Path

log = logging.getLogger('Extract')

# Read directly. Anything textual, plus the formats with a real parser.
TEXT_EXT = {
    '.txt', '.md', '.markdown', '.rst', '.log', '.csv', '.tsv', '.json',
    '.yaml', '.yml', '.xml', '.html', '.htm', '.ini', '.cfg', '.toml', '.env',
    '.py', '.js', '.ts', '.tsx', '.jsx', '.dart', '.java', '.c', '.h', '.cpp',
    '.hpp', '.cs', '.go', '.rs', '.rb', '.php', '.sh', '.ps1', '.bat', '.sql',
    '.tex', '.srt', '.vtt',
}

PDF_EXT = {'.pdf'}
DOCX_EXT = {'.docx'}
IMAGE_EXT = {'.png', '.jpg', '.jpeg', '.bmp', '.tif', '.tiff', '.webp', '.gif'}

# Never opened. Reading these costs time and yields nothing searchable.
SKIP_EXT = {
    '.exe', '.dll', '.sys', '.msi', '.bin', '.iso', '.img', '.pdb', '.obj',
    '.o', '.a', '.lib', '.so', '.dylib', '.class', '.jar', '.pyc', '.pyd',
    '.zip', '.rar', '.7z', '.gz', '.tar', '.xz', '.cab',
    '.mp4', '.mkv', '.avi', '.mov', '.wmv', '.flv', '.webm',
    '.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a',
    '.ttf', '.otf', '.woff', '.woff2', '.eot',
    '.db', '.sqlite', '.sqlite3', '.mdb', '.dat', '.pack', '.idx',
}

# Ceilings, per route.
#
# Text is cheap so the limit is generous. OCR is seconds per image, so a
# forty-megabyte photo is not worth the wall clock during a first crawl.
MAX_TEXT_BYTES = 8 * 1024 * 1024
MAX_OCR_BYTES = 12 * 1024 * 1024
MAX_PDF_PAGES = 30

# What is kept per file. Enough to find it and to show a line of context;
# storing whole documents would make the index bigger than the disk it indexes.
MAX_BODY_CHARS = 60_000


# The three control characters that carry meaning in running text.
KEEP = chr(10) + chr(13) + chr(9)


def clean(text: str) -> str:
    """
    Text safe to put in the index.

    A real Aadhaar PDF on this machine extracts 986 characters, nineteen of
    which are NUL — glyphs the document's font encoding could not map. SQLite
    stops at the first one: those 986 characters were stored as **4**, and the
    only reason a search for "aadhaar" still found the file was that the word
    happened to appear before the damage.

    Everything below a space goes, except the three that carry meaning in
    running text. The replacement is a space rather than nothing, so two words
    either side of a bad glyph do not silently become one.
    """
    if not text:
        return ''
    return ''.join(
        character if character >= ' ' or character in KEEP else ' '
        for character in text
    )


def extract(path: Path, size: int) -> tuple:
    """
    Returns `(body, kind)` — the searchable text and how it was obtained.

    Never raises. A file that cannot be read is recorded as 'failed' and stays
    findable by its name, which is strictly better than an index that stops on
    the first PDF with a broken xref table.
    """
    ext = path.suffix.lower()

    if ext in SKIP_EXT:
        return '', 'skipped'

    try:
        if ext in TEXT_EXT:
            if size > MAX_TEXT_BYTES:
                return '', 'skipped'
            return clean(_read_text(path)), 'text'

        if ext in PDF_EXT:
            body = clean(_read_pdf(path))
            if body.strip():
                return body, 'text'
            # A PDF with no text layer is a scan. That is the interesting case.
            if size <= MAX_OCR_BYTES:
                ocr = clean(_ocr_pdf(path))
                if ocr.strip():
                    return ocr, 'ocr'
            return '', 'skipped'

        if ext in DOCX_EXT:
            return clean(_read_docx(path)), 'text'

        if ext in IMAGE_EXT:
            if size > MAX_OCR_BYTES:
                return '', 'skipped'
            body = clean(_ocr_image(path))
            return (body, 'ocr') if body.strip() else ('', 'skipped')

    except Exception as exc:  # noqa: BLE001 - one unreadable file is not a crawl failure
        log.debug('could not read %s: %s', path, exc)
        return '', 'failed'

    return '', 'skipped'


def _read_text(path: Path) -> str:
    with open(path, 'r', encoding='utf-8', errors='replace') as handle:
        return handle.read(MAX_BODY_CHARS)


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages[:MAX_PDF_PAGES]:
        parts.append(page.extract_text() or '')
        if sum(len(p) for p in parts) > MAX_BODY_CHARS:
            break
    return '\n'.join(parts)[:MAX_BODY_CHARS]


def _read_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(' '.join(cell.text for cell in row.cells))
    return '\n'.join(parts)[:MAX_BODY_CHARS]


# Windows' own OCR, through PowerShell.
#
# WinRT from Python needs the `winsdk` package, which is not installed and is a
# heavier dependency than one PowerShell call per image at index time.
#
# The path travels in an environment variable, not as an argument and not
# interpolated into the script. `-Command` does not bind a `param()` block —
# that is `-File` behaviour — so the first version silently received an empty
# path and failed inside GetFileFromPathAsync on a null. An environment
# variable also means a filename containing a quote or a dollar sign cannot
# become part of the script.
_OCR_PS = r'''
$Path = $env:DEX_OCR_PATH
Add-Type -AssemblyName System.Runtime.WindowsRuntime | Out-Null
$null = [Windows.Media.Ocr.OcrEngine, Windows.Media, ContentType=WindowsRuntime]
$null = [Windows.Graphics.Imaging.BitmapDecoder, Windows.Graphics.Imaging, ContentType=WindowsRuntime]
$null = [Windows.Storage.StorageFile, Windows.Storage, ContentType=WindowsRuntime]

$await = [System.WindowsRuntimeSystemExtensions].GetMethods() | Where-Object {
  $_.Name -eq 'GetAwaiter' -and $_.GetParameters().Count -eq 1 -and
  $_.GetParameters()[0].ParameterType.Name -eq 'IAsyncOperation`1'
} | Select-Object -First 1

function Wait-Async($op, $type) {
  $m = $await.MakeGenericMethod($type)
  $task = $m.Invoke($null, @($op))
  $task.GetResult()
}

$engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromUserProfileLanguages()
if (-not $engine) { exit 2 }

$file = Wait-Async ([Windows.Storage.StorageFile]::GetFileFromPathAsync($Path)) ([Windows.Storage.StorageFile])
$stream = Wait-Async ($file.OpenAsync([Windows.Storage.FileAccessMode]::Read)) ([Windows.Storage.Streams.IRandomAccessStream])
$decoder = Wait-Async ([Windows.Graphics.Imaging.BitmapDecoder]::CreateAsync($stream)) ([Windows.Graphics.Imaging.BitmapDecoder])
$bitmap = Wait-Async ($decoder.GetSoftwareBitmapAsync()) ([Windows.Graphics.Imaging.SoftwareBitmap])
$result = Wait-Async ($engine.RecognizeAsync($bitmap)) ([Windows.Media.Ocr.OcrResult])
[Console]::Out.Write($result.Text)
'''

OCR_TIMEOUT = 45


def _ocr_image(path: Path) -> str:
    """Words printed inside an image. Empty when there are none."""
    try:
        result = subprocess.run(
            ['powershell', '-NoProfile', '-NonInteractive',
             '-ExecutionPolicy', 'Bypass', '-Command', _OCR_PS],
            capture_output=True, text=True, timeout=OCR_TIMEOUT,
            creationflags=0x08000000 if os.name == 'nt' else 0,
            env={**os.environ, 'DEX_OCR_PATH': str(path.resolve())},
        )
    except Exception:  # noqa: BLE001 - a timeout is a skip, not a crawl failure
        return ''
    return (result.stdout or '')[:MAX_BODY_CHARS] if result.returncode == 0 else ''


def _ocr_pdf(path: Path) -> str:
    """
    A scanned PDF, page by page, through the image path.

    Only the first few pages. An identity document is page one, and OCR'ing a
    three-hundred-page scanned book during a first crawl would be a large cost
    for a small return.
    """
    try:
        import tempfile

        from pypdf import PdfReader
    except ImportError:
        return ''

    parts = []
    try:
        reader = PdfReader(str(path))
        for page in reader.pages[:3]:
            for image in list(getattr(page, 'images', []))[:2]:
                with tempfile.NamedTemporaryFile(
                    suffix=Path(image.name).suffix or '.png', delete=False,
                ) as handle:
                    handle.write(image.data)
                    temporary = handle.name
                try:
                    parts.append(_ocr_image(Path(temporary)))
                finally:
                    try:
                        os.unlink(temporary)
                    except OSError:
                        pass
            if sum(len(p) for p in parts) > MAX_BODY_CHARS:
                break
    except Exception:  # noqa: BLE001
        return '\n'.join(parts)

    return '\n'.join(parts)[:MAX_BODY_CHARS]


def ocr_available() -> bool:
    """Whether Windows OCR can be used. Asked once, reported in the stats."""
    try:
        result = subprocess.run(
            ['powershell', '-NoProfile', '-NonInteractive', '-Command',
             '$null = [Windows.Media.Ocr.OcrEngine, Windows.Media, ContentType=WindowsRuntime]; '
             'if ([Windows.Media.Ocr.OcrEngine]::TryCreateFromUserProfileLanguages()) '
             '{ "yes" } else { "no" }'],
            capture_output=True, text=True, timeout=30,
            creationflags=0x08000000 if os.name == 'nt' else 0,
        )
        return 'yes' in (result.stdout or '')
    except Exception:  # noqa: BLE001
        return False
